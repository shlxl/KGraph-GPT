import os
import time
import re
import io
import zipfile
import streamlit as st
from typing import List, Optional, Dict, Any
from pathlib import Path
from pydantic import BaseModel, Field
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from pyvis.network import Network
import json
import streamlit.components.v1 as components
from dotenv import load_dotenv
from src.parsers.markdown_parser import MarkdownMultiDocumentParser

# Import parsers for different file types
from PyPDF2 import PdfReader
import docx
from odf.opendocument import load as load_odt
from odf import text as odf_text, teletype as odf_teletype
from bs4 import BeautifulSoup
from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound
from neo4j import GraphDatabase

# --- CONFIGURATION ---

__version__ = "0.1.0"

st.set_page_config(page_title="文本知识图谱提取器")

# Load environment variables
load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
NEO4J_URI = os.getenv("NEO4J_URI")
NEO4J_USER = os.getenv("NEO4J_USER")
NEO4J_PASSWORD = os.getenv("NEO4J_PASSWORD")

SUPPORTED_FILE_EXTENSIONS = {".txt", ".pdf", ".docx", ".md", ".html", ".htm", ".odt"}

# Load REL_SET configurations
REL_SETS = {}
try:
    with open("GraphRAG-RELSET-AI-RAG-zh.json", "r", encoding="utf-8") as f:
        REL_SETS["GraphRAG-RELSET-AI-RAG-zh"] = json.load(f)
    with open("GraphRAG-RELSET-GenericWeb-zh.json", "r", encoding="utf-8") as f:
        REL_SETS["GraphRAG-RELSET-GenericWeb-zh"] = json.load(f)
except FileNotFoundError:
    st.error("REL_SET JSON files not found. Please ensure 'GraphRAG-RELSET-AI-RAG-zh.json' and 'GraphRAG-RELSET-GenericWeb-zh.json' are in the project root.")
    st.stop()
except json.JSONDecodeError:
    st.error("Error decoding REL_SET JSON files. Please check their format.")
    st.stop()

# Preload available mentions.jsonl files for entity normalization
MENTIONS_DIR = Path("GraphRAG-Extract-Best-Example-CoralWind-zh/gold")
AVAILABLE_MENTIONS_FILES = {}
if MENTIONS_DIR.exists():
    for mentions_file in sorted(MENTIONS_DIR.glob("*.jsonl")):
        AVAILABLE_MENTIONS_FILES[mentions_file.name] = mentions_file

EXAMPLE_DIRECTORIES = {}
EXAMPLE_BASE = Path("GraphRAG-Extract-Best-Example-CoralWind-zh/corpus")
if EXAMPLE_BASE.exists():
    EXAMPLE_DIRECTORIES["CoralWind 官方示例语料 (8 篇 TXT)"] = EXAMPLE_BASE

# Load prompt template
try:
    with open("GraphRAG_prompt.md", "r", encoding="utf-8") as f:
        PROMPT_TEMPLATE = f.read()
except FileNotFoundError:
    st.error("Prompt template 'GraphRAG_prompt.md' not found. Please ensure it is in the project root.")
    st.stop()

# --- Pydantic Models for Graph Structure ---

class Node(BaseModel):
    id: str = Field(..., description="Unique identifier for the node.")
    type: str = Field("Unknown", description="The type or label of the node.")
    properties: Optional[dict] = Field(None, description="Additional properties of the node.")
    color: Optional[str] = Field(None, description="The color of the node.")

class Relationship(BaseModel):
    source: Node
    target: Node
    type: str
    properties: Optional[dict] = Field(None, description="Additional properties of the relationship.")
    qualifiers: Optional[dict] = Field(None, description="Additional qualifiers for the relationship, such as time, amount, or role.")
    color: Optional[str] = Field(None, description="The color of the relationship.")
    evidence: Optional[List[Dict[str, Any]]] = Field(None, description="List of evidence supporting the relationship, including document ID and sentence IDs.")
    confidence: Optional[float] = Field(None, description="Confidence score (0.0-1.0) for the extracted relationship.")

class Metadata(BaseModel):
    source: str = Field(..., description="The source of the data, such as a file name or URL.")
    timestamp: str = Field(..., description="The timestamp of the graph creation.")
    doc_id: Optional[str] = Field(None, description="The ID of the document from which the graph was extracted.")
    doc_date: Optional[str] = Field(None, description="The date of the document from which the graph was extracted.")

class KnowledgeGraph(BaseModel):
    nodes: List[Node]
    relationships: List[Relationship]
    metadata: Optional[Metadata] = None


PrimitiveTypes = (str, int, float, bool)

def sanitize_property_value(value):
    if value is None:
        return None
    if isinstance(value, PrimitiveTypes):
        return value
    if isinstance(value, list):
        primitives_only = []
        for item in value:
            if item is None:
                continue
            if isinstance(item, PrimitiveTypes):
                primitives_only.append(item)
            else:
                return json.dumps(value, ensure_ascii=False)
        return primitives_only
    return json.dumps(value, ensure_ascii=False)


def sanitize_properties(properties: Optional[Dict[str, Any]]) -> Dict[str, Any]:
    sanitized = {}
    if not properties:
        return sanitized
    for key, value in properties.items():
        sanitized_value = sanitize_property_value(value)
        if sanitized_value is not None:
            sanitized[key] = sanitized_value
    return sanitized


def sanitize_relationship_type(rel_type: Optional[str]) -> str:
    if not rel_type:
        return "RELATIONSHIP"
    cleaned = re.sub(r"[^A-Za-z0-9_]", "_", rel_type)
    cleaned = re.sub(r"_+", "_", cleaned).strip("_")
    return cleaned.upper() if cleaned else "RELATIONSHIP"


def process_markdown_content(markdown_content: str, fallback_doc_id: str, fallback_source: str) -> List[Dict[str, Any]]:
    documents = []
    if "/corpus/" in markdown_content:
        parser = MarkdownMultiDocumentParser()
        parsed_docs = parser.parse(markdown_content)
        if parsed_docs:
            documents.extend(parsed_docs)
        else:
            st.warning(f"未能在 {fallback_source} 中找到任何可解析的多文档内容。")
    else:
        cleaned_text = re.sub(r'[\*\#\`\>]', '', markdown_content)
        documents.append({
            "doc_id": fallback_doc_id,
            "source": fallback_source,
            "date": time.strftime("%Y-%m-%d"),
            "text_with_sentence_ids": cleaned_text
        })
    return documents


def get_text_from_path(file_path: Path) -> str:
    file_extension = file_path.suffix.lower()
    try:
        if file_extension == ".pdf":
            text = ""
            with file_path.open("rb") as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
            return text
        elif file_extension == ".docx":
            with file_path.open("rb") as f:
                doc = docx.Document(f)
            text = "\n".join(para.text for para in doc.paragraphs)
            return text
        elif file_extension == ".odt":
            doc = load_odt(str(file_path))
            all_paras = doc.getElementsByType(odf_text.P)
            text = "\n".join(odf_teletype.extractText(para) for para in all_paras)
            return text
        elif file_extension in [".html", ".htm"]:
            with file_path.open("r", encoding="utf-8") as f:
                return BeautifulSoup(f.read(), "html.parser").get_text()
        elif file_extension == ".txt":
            return file_path.read_text(encoding="utf-8")
        else:
            st.warning(f"目录中文件 {file_path.name} 的格式暂不支持。")
            return ""
    except Exception as e:
        st.error(f"读取文件 {file_path} 时发生错误: {e}")
        return ""


def load_documents_from_directory(directory_path: Path) -> List[Dict[str, Any]]:
    documents = []
    if not directory_path.exists() or not directory_path.is_dir():
        st.error(f"目录 {directory_path} 不存在或不是有效目录。")
        return documents

    matched_files = sorted([p for p in directory_path.rglob("*") if p.suffix.lower() in SUPPORTED_FILE_EXTENSIONS and p.is_file()])

    if not matched_files:
        st.warning(f"目录 {directory_path} 中未找到支持的文件类型: {', '.join(sorted(SUPPORTED_FILE_EXTENSIONS))}")
        return documents

    for file_path in matched_files:
        extension = file_path.suffix.lower()
        relative_name = str(file_path.relative_to(directory_path))
        if extension == ".md":
            markdown_content = file_path.read_text(encoding="utf-8")
            documents.extend(process_markdown_content(markdown_content, relative_name, relative_name))
        else:
            text = get_text_from_path(file_path)
            documents.append({
                "doc_id": relative_name,
                "source": relative_name,
                "date": time.strftime("%Y-%m-%d"),
                "text_with_sentence_ids": text
            })
    return documents

# --- DATABASE LOGIC ---

class Neo4jDatabase:
    def __init__(self, uri, user, password):
        self._driver = GraphDatabase.driver(uri, auth=(user, password))

    def close(self):
        self._driver.close()

    def save_graph(self, graph: KnowledgeGraph):
        with self._driver.session() as session:
            metadata_properties = graph.metadata.model_dump() if graph.metadata else {}
            for node in graph.nodes:
                session.execute_write(self._create_node, node, metadata_properties)
            for rel in graph.relationships:
                session.execute_write(self._create_relationship, rel, metadata_properties)

    @staticmethod
    def _create_node(tx, node: Node, metadata: dict):
        node_properties = sanitize_properties(node.properties)
        query = (
            "MERGE (n:Node {id: $id}) "
            "SET n.type = $type "
            "SET n += $properties "
            "SET n.color = $color "
            "SET n.doc_id = $metadata.doc_id "
            "SET n.doc_date = $metadata.doc_date "
            "SET n.source = $metadata.source "
            "SET n.timestamp = $metadata.timestamp"
        )
        tx.run(query, id=node.id, type=node.type, properties=node_properties, color=node.color, metadata=metadata)

    @staticmethod
    def _create_relationship(tx, rel: Relationship, metadata: dict):
        rel_properties = sanitize_properties(rel.properties)
        evidence_value = sanitize_property_value(rel.evidence)
        sanitized_rel_type = sanitize_relationship_type(rel.type)
        query = (
            "MATCH (a:Node) WHERE a.id = $source_id "
            "MATCH (b:Node) WHERE b.id = $target_id "
            f"MERGE (a)-[r:{sanitized_rel_type}]->(b) "
            "SET r += $properties "
            "SET r.color = $color "
            "SET r.evidence = $evidence "
            "SET r.confidence = $confidence "
            "SET r.doc_id = $metadata.doc_id "
            "SET r.doc_date = $metadata.doc_date "
            "SET r.source = $metadata.source "
            "SET r.timestamp = $metadata.timestamp"
        )
        tx.run(
            query,
            source_id=rel.source.id,
            target_id=rel.target.id,
            properties=rel_properties,
            color=rel.color,
            evidence=evidence_value,
            confidence=rel.confidence,
            metadata=metadata,
        )

# --- MODEL INITIALIZATION ---

def get_llm(model_name: str):
    return ChatGoogleGenerativeAI(model=model_name, temperature=0, google_api_key=GOOGLE_API_KEY)

# --- CORE LOGIC ---

# (File and YouTube parsing functions remain the same...)

def is_youtube_url(url: str) -> bool:
    youtube_regex = (
        r'(https?://)?(www\.)?'
        r'(youtube|youtu|youtube-nocookie)\.(com|be)/'
        r'(watch\?v=|embed/|v/|.+\?v=)?([^&=%\?]{11})')
    return re.match(youtube_regex, url) is not None

def get_text_from_youtube(url: str) -> str:
    try:
        video_id_match = re.search(r"(?:v=|\/|embed\/|v\/|youtu\.be\/)([a-zA-Z0-9_-]{11})", url)
        if not video_id_match:
            st.error("无法从URL中提取有效的YouTube视频ID。")
            return ""
        video_id = video_id_match.group(1)
        api = YouTubeTranscriptApi()
        transcript_list = api.list(video_id)
        try:
            transcript = transcript_list.find_transcript(['zh-Hans', 'zh-Hant', 'en'])
        except NoTranscriptFound:
            try:
                transcript = transcript_list.find_generated_transcript(['zh-Hans', 'zh-Hant', 'en'])
            except NoTranscriptFound:
                st.error(f"视频 {video_id} 未找到中文或英文字幕。")
                return ""
        full_transcript = " ".join([item.text for item in transcript.fetch()])
        return full_transcript
    except Exception as e:
        st.error(f"获取YouTube字幕时发生错误: {e}")
        return ""

def get_text_from_file(uploaded_file) -> str:
    try:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        text = ""
        if file_extension == ".pdf":
            reader = PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        elif file_extension == ".docx":
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        elif file_extension == ".odt":
            doc = load_odt(uploaded_file)
            all_paras = doc.getElementsByType(odf_text.P)
            for para in all_paras:
                text += odf_teletype.extractText(para) + "\n"
            return text
        elif file_extension in [".html", ".htm"]:
            return BeautifulSoup(uploaded_file.read(), "html.parser").get_text()
        elif file_extension == ".md":
            md_text = uploaded_file.read().decode("utf-8")
            text = re.sub(r'[\*\#\`\>]', '', md_text)
            return text
        elif file_extension == ".txt":
            return uploaded_file.read().decode("utf-8")
        else:
            st.warning(f"不支持的文件格式: {file_extension}")
            return ""
    except Exception as e:
        st.error(f"解析文件时发生错误: {e}")
        return ""

def generate_graph(text: str, source: str, model_name: str, node_color: str, edge_color: str, rel_set_name: str, doc_id: Optional[str] = None, doc_date: Optional[str] = None) -> KnowledgeGraph:
    llm = get_llm(model_name)
    structured_llm = llm.with_structured_output(KnowledgeGraph)

    selected_rel_set = REL_SETS.get(rel_set_name)
    if not selected_rel_set:
        raise ValueError(f"REL_SET '{rel_set_name}' not found.")

    # Format relations for the prompt
    relations_list = [json.dumps(rel["name"]) for rel in selected_rel_set["relations"]]
    formatted_relations = ",\n  ".join(relations_list)
    
    # Format qualifiers for the prompt
    formatted_qualifiers = json.dumps(selected_rel_set["qualifiers"], indent=2, ensure_ascii=False)

    # Format policies for the prompt
    formatted_policies = json.dumps(selected_rel_set["policies"], indent=2, ensure_ascii=False)
    formatted_policies = formatted_policies.replace("{", "{{").replace("}", "}}")

    # Format authority order for the prompt
    formatted_authority_order = json.dumps(selected_rel_set["authority_order_for_disputes"], indent=2, ensure_ascii=False)
    formatted_authority_order = formatted_authority_order.replace("{", "{{").replace("}", "}}")

    # Construct the system prompt using the template and REL_SET data
    system_prompt_content = PROMPT_TEMPLATE.replace("{REL_SET}", formatted_relations).replace("{QUALIFIERS}", formatted_qualifiers)
    system_prompt_content = system_prompt_content.replace("{{POLICIES}}", formatted_policies).replace("{{AUTHORITY_ORDER}}", formatted_authority_order)
    
    # Populate document-specific metadata in the prompt
    system_prompt_content = system_prompt_content.replace("{{doc_id}}", doc_id if doc_id else "未知")
    system_prompt_content = system_prompt_content.replace("{{doc_date}}", doc_date if doc_date else "未知")
    system_prompt_content = system_prompt_content.replace("{{source_name}}", source)
    
    system_prompt = """从文本中提取知识图谱。请识别出所有的实体作为节点，以及它们之间的关系。
    确保节点具有唯一的ID（通常是实体的名称）和类型（例如：人、地点、组织、概念）。
    如果实体或关系有额外的属性（例如日期、数量、职位、事件描述等），请将它们提取到'properties'字段中。
    特别地，如果关系是双向的（例如‘合作’、‘同事’、‘配偶’），请为每个方向都生成一条关系边（例如 A-合作->B 和 B-合作->A）。
    你必须严格按照KnowledgeGraph的格式返回结果。
    """

    prompt = ChatPromptTemplate.from_messages([
        ("system", system_prompt_content),
        ("human", "请从以下文本中提取知识图谱：\n\n{text}")
    ])
    
    chain = prompt | structured_llm
    graph = chain.invoke({"text": text})
    graph.metadata = Metadata(source=source, timestamp=time.strftime("%Y-%m-%d %H:%M:%S"), doc_id=doc_id, doc_date=doc_date)
    
    for node in graph.nodes:
        node.color = node_color
    for relationship in graph.relationships:
        relationship.color = edge_color
    
    return graph


# --- UI & VISUALIZATION ---

st.title("文本知识图谱提取器")
st.caption(f"Version {__version__}")
st.caption("从文本中提取知识图谱并进行可视化。")

text_input = st.text_area("粘贴文本（或输入YouTube链接）:", key="text_input")
uploaded_file = st.file_uploader(
    "或者上传一个文件", 
    type=['txt', 'pdf', 'docx', 'md', 'html', 'htm', 'odt'],
    key="uploaded_file"
)

directory_path_input = st.text_input(
    "输入本地目录路径以批量处理（可选）",
    value="",
    placeholder="例如：GraphRAG-Extract-Best-Example-CoralWind-zh/corpus"
)

example_directory_selection = []
if EXAMPLE_DIRECTORIES:
    example_directory_selection = st.multiselect(
        "或直接选择内置示例目录（可多选）",
        list(EXAMPLE_DIRECTORIES.keys())
    )

mention_selection_options = ["不使用实体规范化"] + list(AVAILABLE_MENTIONS_FILES.keys())
selected_mentions_option = st.selectbox(
    "选择实体规范化数据源 (提取自 GraphRAG-Extract-Best-Example-CoralWind-zh/gold)",
    mention_selection_options
)
selected_mentions_path = AVAILABLE_MENTIONS_FILES.get(selected_mentions_option)

model_selection = st.selectbox(
    "选择一个模型:",
    ("gemini-2.5-pro", "gemini-1.5-pro", "gemini-1.5-flash")
)

rel_set_selection = st.selectbox(
    "选择关系集配置:",
    list(REL_SETS.keys())
)

with st.expander("自定义颜色"):
    node_color = st.color_picker("选择节点颜色", "#FFADAD")
    edge_color = st.color_picker("选择边颜色", "#9BF6FF")
    node_shape = st.selectbox(
        "选择节点形状:",
        ("dot", "square", "triangle", "star")
    )
    node_size = st.slider("选择节点大小", 1, 100, 25)
    edge_width = st.slider("选择边大小", 1, 20, 5)
    physics_enabled = st.checkbox("启用物理效果", value=True)
    edge_arrow_style = st.selectbox(
        "选择边箭头样式:",
        ("arrow", "circle", "box", "curve")
    )
    font_selection = st.selectbox(
        "选择字体:",
        ("Arial", "Verdana", "Tahoma", "Georgia")
    )
    background_color = st.color_picker("选择背景颜色", "#222222")
    font_color = st.color_picker("选择字体颜色", "#FFFFFF")
    node_highlight_color = st.color_picker("选择节点高亮颜色", "#FFD700")
    edge_highlight_color = st.color_picker("选择边高亮颜色", "#FFD700")
    edge_hover_color = st.color_picker("选择边悬停颜色", "#FF4500")
    node_border_width = st.slider("选择节点边框宽度", 1, 10, 1)
    node_shadow = st.checkbox("启用节点阴影", value=True)
    edge_smoothness = st.selectbox(
        "选择边平滑度:",
        ("dynamic", "continuous", "discrete", "cubicBezier")
    )
    zoom_enabled = st.checkbox("启用缩放", value=True)
    drag_enabled = st.checkbox("启用拖动", value=True)
    hover_enabled = st.checkbox("启用悬停效果", value=True)

with st.expander("自定义物理效果"):
    st.write("仅当'启用物理效果'勾选时生效")
    physics_solver = st.selectbox("选择物理求解器", ("barnesHut", "repulsion", "hierarchicalRepulsion"))
    physics_gravity = st.slider("重力", -1000, 0, -200)
    physics_friction = st.slider("摩擦力", 0.0, 1.0, 0.1, 0.01)
    physics_repulsion = st.slider("斥力", 0, 5000, 1000)

with st.expander("自定义聚类设置"):
    cluster_enabled = st.checkbox("启用聚类", value=False)
    cluster_size = st.slider("聚类大小", 1, 100, 10)
    cluster_iterations = st.slider("聚类迭代次数", 1, 100, 10)

with st.expander("自定义布局设置"):
    st.write("仅当'分层布局'勾选时生效")
    hierarchical_sort_method = st.selectbox("分层排序方法", ("hubsize", "directed", "undirected"))
    hierarchical_level_separation = st.slider("分层级别间距", 50, 500, 150)
    hierarchical_node_spacing = st.slider("分层节点间距", 50, 500, 100)
    hierarchical_tree_spacing = st.slider("分层树间距", 50, 500, 200)

with st.expander("自定义物理稳定设置"):
    stabilization_enabled = st.checkbox("启用物理稳定", value=True)
    stabilization_iterations = st.slider("稳定迭代次数", 50, 2000, 1000)
    stabilization_fit = st.checkbox("稳定后适应视图", value=True)

with st.expander("自定义操作设置"):
    manipulation_enabled = st.checkbox("启用操作", value=False)
    manipulation_add_node = st.checkbox("允许添加节点", value=False)
    manipulation_add_edge = st.checkbox("允许添加边", value=False)
    manipulation_edit_node = st.checkbox("允许编辑节点", value=False)
    manipulation_edit_edge = st.checkbox("允许编辑边", value=False)

with st.expander("自定义导航按钮"):
    navigation_buttons_enabled = st.checkbox("启用导航按钮", value=False)

with st.expander("自定义工具提示设置"):
    node_tooltip_enabled = st.checkbox("启用节点工具提示", value=True)
    edge_tooltip_enabled = st.checkbox("启用边工具提示", value=True)

layout_selection = st.selectbox(
    "选择一个布局:",
    ("Hierarchical", "Force Atlas 2")
)

generate_button = st.button("生成图谱")
clear_button = st.button("清除")

if clear_button:
    st.session_state.text_input = ""
    st.session_state.uploaded_file = None
    st.rerun()

def normalize_entities(graph: KnowledgeGraph, mentions_data: List[Dict[str, Any]]) -> KnowledgeGraph:
    """
    规范化知识图谱中的实体ID，根据提供的mentions数据进行别名映射。
    合并具有相同规范化ID的节点属性，并确保关系源和目标引用实际的规范化Node对象。
    """
    id_mapping = {}
    for idx, mention_entry in enumerate(mentions_data, start=1):
        canonical_id = mention_entry.get("canonical_id")
        name = mention_entry.get("name")
        aliases = mention_entry.get("aliases", [])

        if not canonical_id or not name:
            st.warning(f"mentions.jsonl 第 {idx} 行缺少 canonical_id 或 name，已跳过该条记录。")
            continue

        if not isinstance(aliases, list):
            aliases = [aliases]

        # Map the canonical name itself
        id_mapping[name] = canonical_id
        # Map all aliases
        for alias in aliases:
            if alias:
                id_mapping[alias] = canonical_id
    
    # Use a dictionary to store unique normalized nodes, keyed by their normalized ID
    unique_normalized_nodes: Dict[str, Node] = {}
    
    # First pass: Normalize nodes and merge properties
    for node in graph.nodes:
        original_id = node.id
        normalized_id = id_mapping.get(original_id, original_id)
        
        if normalized_id in unique_normalized_nodes:
            # Node with this normalized ID already exists, merge properties
            existing_node = unique_normalized_nodes[normalized_id]
            if node.properties:
                if existing_node.properties:
                    existing_node.properties.update(node.properties)
                else:
                    existing_node.properties = node.properties
            # Prioritize existing node's type if current node's type is "Unknown"
            if existing_node.type == "Unknown" and node.type != "Unknown":
                existing_node.type = node.type
            # Keep existing node's color if current node's color is None
            if existing_node.color is None and node.color is not None:
                existing_node.color = node.color
        else:
            # New normalized node
            node.id = normalized_id
            unique_normalized_nodes[normalized_id] = node

    normalized_nodes_list = list(unique_normalized_nodes.values())

    # Second pass: Normalize relationships, ensuring source/target refer to the correct Node objects
    normalized_relationships = []
    for rel in graph.relationships:
        original_source_id = rel.source.id
        original_target_id = rel.target.id
        
        normalized_source_id = id_mapping.get(original_source_id, original_source_id)
        normalized_target_id = id_mapping.get(original_target_id, original_target_id)
        
        # Ensure source and target nodes exist in our unique_normalized_nodes map
        # If a relationship points to a node that wasn't extracted or normalized,
        # we might need to create a placeholder node or handle it as an error.
        # For now, we assume all related nodes will be present after the first pass.
        if normalized_source_id in unique_normalized_nodes and normalized_target_id in unique_normalized_nodes:
            rel.source = unique_normalized_nodes[normalized_source_id]
            rel.target = unique_normalized_nodes[normalized_target_id]
            normalized_relationships.append(rel)
        else:
            st.warning(f"关系 {rel.type} ({original_source_id} -> {original_target_id}) 引用了未找到的规范化节点，跳过此关系。")

    return KnowledgeGraph(nodes=normalized_nodes_list, relationships=normalized_relationships, metadata=graph.metadata)


if generate_button:
    if not GOOGLE_API_KEY:
        st.error("未找到 GOOGLE_API_KEY。请确保您的 .env 文件已正确设置。")
        st.stop()
    
    documents_to_process = []
    
    if text_input:
        input_text = text_input.strip()
        if input_text:
            if is_youtube_url(input_text):
                documents_to_process.append({
                    "doc_id": "youtube_video",
                    "source": input_text,
                    "date": time.strftime("%Y-%m-%d"),
                    "text_with_sentence_ids": get_text_from_youtube(input_text)
                })
            else:
                documents_to_process.append({
                    "doc_id": "text_input",
                    "source": "用户输入",
                    "date": time.strftime("%Y-%m-%d"),
                    "text_with_sentence_ids": input_text
                })
    
    if uploaded_file:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        if file_extension == ".md":
            markdown_content = uploaded_file.read().decode("utf-8")
            documents_to_process.extend(process_markdown_content(markdown_content, uploaded_file.name, uploaded_file.name))
        else:
            uploaded_file.seek(0)
            documents_to_process.append({
                "doc_id": uploaded_file.name,
                "source": uploaded_file.name,
                "date": time.strftime("%Y-%m-%d"),
                "text_with_sentence_ids": get_text_from_file(uploaded_file)
            })
    
    if directory_path_input.strip():
        directory_path = Path(directory_path_input.strip()).expanduser()
        directory_documents = load_documents_from_directory(directory_path)
        documents_to_process.extend(directory_documents)
    
    if example_directory_selection:
        for example_label in example_directory_selection:
            example_path = EXAMPLE_DIRECTORIES.get(example_label)
            if example_path:
                documents_to_process.extend(load_documents_from_directory(example_path))

    if not documents_to_process or not any(doc["text_with_sentence_ids"] for doc in documents_to_process):
        st.warning("请输入文本、YouTube链接、上传文件，或提供有效的目录路径。")
        st.stop()
    else:
        progress_bar = st.progress(0, text="正在初始化...")
        all_graphs = []
        total_docs = len(documents_to_process)
        doc_source_summary = [doc.get("source") for doc in documents_to_process if doc.get("source")]

        for i, doc_data in enumerate(documents_to_process):
            doc_id = doc_data["doc_id"]
            source = doc_data["source"]
            date = doc_data["date"]
            text_to_process = doc_data["text_with_sentence_ids"]

            if not text_to_process:
                st.info(f"文档 '{doc_id}' 内容为空，跳过处理。")
                continue

            progress_text = f"正在处理文档 {i+1}/{total_docs} ({doc_id})... 调用大语言模型 (这可能需要一些时间)"
            progress_bar.progress((i * 100 // total_docs) + 10, text=progress_text)
            
            try:
                graph = generate_graph(
                    text=text_to_process,
                    source=source,
                    model_name=model_selection,
                    node_color=node_color,
                    edge_color=edge_color,
                    rel_set_name=rel_set_selection,
                    doc_id=doc_id,
                    doc_date=date
                )
                all_graphs.append(graph)
            except Exception as e:
                st.error(f"处理文档 '{doc_id}' 时发生错误: {e}")
                continue
        
        if not all_graphs:
            progress_bar.empty()
            st.warning("未能从任何文档中提取出知识图谱。")
            st.stop()

        # Aggregate all graphs into a single KnowledgeGraph object
        aggregated_nodes = []
        aggregated_relationships = []
        for graph in all_graphs:
            aggregated_nodes.extend(graph.nodes)
            aggregated_relationships.extend(graph.relationships)
        
        # Create a new KnowledgeGraph for the aggregated result
        # The metadata for the aggregated graph can be a summary or the first doc's metadata
        aggregated_graph = KnowledgeGraph(
            nodes=aggregated_nodes,
            relationships=aggregated_relationships,
            metadata=Metadata(source="聚合图谱", timestamp=time.strftime("%Y-%m-%d %H:%M:%S"))
        )

        # --- Entity Normalization/Alias Handling ---
        if selected_mentions_path:
            progress_bar.progress(70, text="正在进行实体规范化...")
            try:
                mentions_data = []
                with open(selected_mentions_path, "r", encoding="utf-8") as mentions_file:
                    for line_number, raw_line in enumerate(mentions_file, start=1):
                        line = raw_line.strip()
                        if not line:
                            continue
                        try:
                            record = json.loads(line)
                        except json.JSONDecodeError as json_err:
                            st.warning(f"{selected_mentions_option} 第 {line_number} 行解析失败: {json_err}")
                            continue

                        missing_fields = [field for field in ("canonical_id", "name") if field not in record]
                        if missing_fields:
                            st.warning(f"{selected_mentions_option} 第 {line_number} 行缺少字段: {', '.join(missing_fields)}，已跳过该条记录。")
                            continue

                        if "aliases" not in record or record["aliases"] is None:
                            record["aliases"] = []
                        elif not isinstance(record["aliases"], list):
                            record["aliases"] = [record["aliases"]]

                        mentions_data.append(record)

                if mentions_data:
                    aggregated_graph = normalize_entities(aggregated_graph, mentions_data)
                    st.success("实体规范化完成。")
                else:
                    st.warning("未找到有效的实体规范化数据，已跳过该步骤。")
            except Exception as e:
                st.error(f"实体规范化失败: {e}")
                st.stop()

        progress_bar.progress(80, text="已获取数据，正在渲染聚合图谱...")
        time.sleep(0.2)

        if aggregated_graph.nodes:
            # Visualize the aggregated graph
            if layout_selection == "Hierarchical":
                net = Network(height="600px", width="100%", bgcolor=background_color, font_color=font_color, notebook=True, directed=True, layout="hierarchical", cdn_resources='in_line')
            else:
                net = Network(height="600px", width="100%", bgcolor=background_color, font_color=font_color, notebook=True, directed=True, cdn_resources='in_line')
                net.force_atlas_2based()

            if navigation_buttons_enabled:
                net.show_buttons()
            
            net.set_options(f"""
            var options = {{
              "nodes": {{
                "borderWidth": {node_border_width},
                "shadow": {str(node_shadow).lower()},
                "font": {{
                  "face": "{font_selection}"
                }},
                "color": {{
                  "highlight": {{
                    "border": "{node_highlight_color}",
                    "background": "{node_highlight_color}"
                  }}
                }},
                "title": {str(node_tooltip_enabled).lower()}
              }},
              "edges": {{
                "smooth": {{
                  "type": "{edge_smoothness}"
                }},
                "font": {{
                  "face": "{font_selection}"
                }},
                "color": {{
                  "highlight": "{edge_highlight_color}",
                  "hover": "{edge_hover_color}"
                }},
                "title": {str(edge_tooltip_enabled).lower()}
              }},
              "interaction": {{
                "zoomView": {str(zoom_enabled).lower()},
                "dragView": {str(drag_enabled).lower()},
                "hover": {str(hover_enabled).lower()},
                "navigationButtons": {str(navigation_buttons_enabled).lower()}
              }},
              "physics": {{
                "enabled": {str(physics_enabled).lower()},
                "solver": "{physics_solver}",
                "{physics_solver}": {{
                  "gravitationalConstant": {physics_gravity},
                  "centralGravity": {physics_friction},
                  "springLength": {physics_repulsion},
                  "springConstant": 0.001,
                  "damping": 0.09
                }},
                "stabilization": {{
                  "enabled": {str(stabilization_enabled).lower()},
                  "iterations": {stabilization_iterations},
                  "fit": {str(stabilization_fit).lower()}
                }}
              }},
              "clustering": {{
                "enabled": {str(cluster_enabled).lower()},
                "clusterEdgeThreshold": {cluster_size},
                "maxEdgeLength": {cluster_iterations}
              }},
              "layout": {{
                "hierarchical": {{
                  "sortMethod": "{hierarchical_sort_method}",
                  "levelSeparation": {hierarchical_level_separation},
                  "nodeSpacing": {hierarchical_node_spacing},
                  "treeSpacing": {hierarchical_tree_spacing}
                }}
              }},
              "manipulation": {{
                "enabled": {str(manipulation_enabled).lower()},
                "addNode": {str(manipulation_add_node).lower()},
                "addEdge": {str(manipulation_add_edge).lower()},
                "editNode": {str(manipulation_edit_node).lower()},
                "editEdge": {str(manipulation_edit_edge).lower()}
              }}
            }}
            """)
            
            for node in aggregated_graph.nodes:
                title = node.model_dump_json(indent=2)
                net.add_node(node.id, label=node.id, title=title, color=node_color, shape=node_shape, size=node_size)
            for edge in aggregated_graph.relationships:
                title = edge.model_dump_json(indent=2)
                net.add_edge(edge.source.id, edge.target.id, label=edge.type, color=edge_color, title=title, width=edge_width, arrows=edge_arrow_style)
            
            graph_html_path = "temp_graph.html"
            net.save_graph(graph_html_path)
            with open(graph_html_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            
            progress_bar.progress(90, text="图谱渲染完毕，正在存入Neo4j数据库...")
            # Save to Neo4j
            if NEO4J_URI and NEO4J_USER and NEO4J_PASSWORD:
                try:
                    db = Neo4jDatabase(NEO4J_URI, NEO4J_USER, NEO4J_PASSWORD)
                    db.save_graph(aggregated_graph)
                    db.close()
                    st.success("聚合图谱已成功存入Neo4j数据库！")
                except Exception as e:
                    st.error(f"连接或写入Neo4j数据库时发生错误: {e}")
            else:
                st.info("未配置Neo4j环境变量，跳过数据库存储。")

            progress_bar.progress(100)
            st.success("聚合图谱生成成功！")
            components.html(html_content, height=620)
            json_payload = aggregated_graph.model_dump()
            download_json = json.dumps(json_payload, ensure_ascii=False, indent=2)
            st.download_button(
                label="下载聚合图谱 (JSON)",
                data=download_json,
                file_name="aggregated_knowledge_graph.json",
                mime="application/json"
            )

            submission_zip = io.BytesIO()
            with zipfile.ZipFile(submission_zip, "w", zipfile.ZIP_DEFLATED) as zip_file:
                zip_file.writestr("aggregated_knowledge_graph.json", download_json)
                zip_file.writestr("graph.html", html_content)
                run_metadata = {
                    "generated_at": time.strftime("%Y-%m-%d %H:%M:%S"),
                    "document_sources": doc_source_summary,
                    "selected_mentions": selected_mentions_option if selected_mentions_path else "未使用",
                    "example_directories": example_directory_selection,
                    "custom_directory": directory_path_input.strip() or "未提供",
                    "model": model_selection,
                    "rel_set": rel_set_selection
                }
                zip_file.writestr("run_metadata.json", json.dumps(run_metadata, ensure_ascii=False, indent=2))
            submission_zip.seek(0)
            st.download_button(
                label="下载提交包 (ZIP，含HTML+JSON+Metadata)",
                data=submission_zip.getvalue(),
                file_name="knowledge_graph_submission.zip",
                mime="application/zip"
            )
            if os.path.exists(graph_html_path):
                os.remove(graph_html_path)
        else:
            progress_bar.progress(100)
            st.warning("未能从聚合文档中提取出任何实体和关系。")
        
        progress_bar.empty()
